import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

def add_heading(text, level=1):
    doc.add_heading(text, level=level)

# Cover Page
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("IUBAT – INTERNATIONAL UNIVERSITY OF BUSINESS AGRICULTURE AND TECHNOLOGY\n")
run.bold = True
run.font.size = Pt(14)
run2 = title.add_run("College of Engineering and Technology (CEAT)\nDepartment of Computer Science & Engineering\n\n")
run2.bold = True
run2.font.size = Pt(12)

run3 = title.add_run("ASSIGNMENT REPORT\n")
run3.bold = True
run3.font.size = Pt(16)

run4 = title.add_run("Autonomous Vehicle Operational Control and Incident Response Simulation\n\n")
run4.bold = True
run4.font.size = Pt(14)

doc.add_paragraph("Student Name: Mostak Shahriar")
doc.add_paragraph("Student ID: 23303076")
doc.add_paragraph("Course: Visual Programming")
doc.add_paragraph("Course Code: CSC 439")
doc.add_paragraph("Semester: Summer 2026")
doc.add_paragraph("Instructor: Sheekar Banerjee, Lecturer, Dept. of CSE, IUBAT\n")
doc.add_paragraph("Submission Type: Assignment Documentation and C# Source Code")
doc.add_paragraph("(Prepared for academic submission)\n")

doc.add_page_break()

# TOC
add_heading("Table of Contents")
toc_items = [
    "1. Introduction",
    "2. System Objectives and Requirements Mapping",
    "3. Program Structure and Architecture",
    "4. Data Integrity and Data Annotation",
    "5. Exception Handling and Incident Response",
    "6. Delegate-Based Event Notification",
    "7. LINQ and Dynamic Vehicle Data Management",
    "8. User-Centric Control Panel (GUI)",
    "9. Operational Safety, Reliability and Design Justification",
    "10. Testing Scenarios and Expected Results",
    "11. Data Persistence (JSON)",
    "12. Conclusion"
]
for item in toc_items:
    doc.add_paragraph(item)

doc.add_page_break()

# Sections
add_heading("1. Introduction")
doc.add_paragraph("The assignment brief asks for a prototype Self-driving Car Simulation System in C# that can manage autonomous vehicle statuses, simulate dynamic driving conditions, respond to system exceptions, and communicate critical operational events to mission control in real time. The intended system is for internal simulation use and must demonstrate operational safety, exception handling, flexible event management, dynamic data operations, and data integrity validation.")

add_heading("2. System Objectives and Requirements Mapping")
doc.add_paragraph("The proposed solution is designed as a modular Windows Forms simulation application. Each vehicle is represented by a Car object. A central in-memory dictionary manages active vehicles, while LINQ is used for filtering, sorting, and summarizing data. A delegate-based alert mechanism sends incident and mission events to a mission-control panel. Data annotations and explicit validation rules protect the integrity of critical vehicle fields.")

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Assignment Requirement'
hdr_cells[1].text = 'Implementation Choice'
hdr_cells[2].text = 'Purpose'
data = [
    ("Car entity definition", "Car class + enums + DataAnnotations", "Represents identity, range, battery, priority, and log data."),
    ("Incident handling", "Custom exceptions + safe-state logic", "Prevents crashes and moves unsafe cars to controlled states."),
    ("Event delegation", "AlertDispatcher + event notifications", "Sends route, obstacle, and restricted-zone alerts."),
    ("Dynamic data management", "FleetManager with LINQ queries", "Provides expressive filtering and real-time fleet summaries."),
    ("User interface", "MainForm with TabControl & DataGridView", "Offers a user-centric dashboard to monitor and command the fleet."),
    ("Data Persistence", "System.Text.Json Serialization", "Saves and loads fleet data automatically to/from fleet.json.")
]
for item in data:
    row_cells = table.add_row().cells
    row_cells[0].text = item[0]
    row_cells[1].text = item[1]
    row_cells[2].text = item[2]

doc.add_paragraph("")

add_heading("3. Program Structure and Architecture")
doc.add_paragraph("The system is divided into clear functional components to ensure separation of concerns:")
doc.add_paragraph("• Models (Car.cs, Enums.cs): Defines the domain entities, enumerations (Status, Priority, Events), and internal operational logs.", style='List Bullet')
doc.add_paragraph("• Business Logic (FleetManager.cs): Manages the fleet collection, exposes LINQ queries, and acts as the central executor for all simulated incidents.", style='List Bullet')
doc.add_paragraph("• Alert System (AlertSystem.cs): Decouples the simulation events from the GUI using C# delegates and events.", style='List Bullet')
doc.add_paragraph("• Incident Hierarchy (Incidents.cs): Defines a common SimulationException base class to safely handle all failures.", style='List Bullet')
doc.add_paragraph("• User Interface (Program.cs): The MainForm acts as the mission control panel, rendering data grids, alerts, and vehicle logs across intuitive tabs.", style='List Bullet')

add_heading("4. Data Integrity and Data Annotation")
doc.add_paragraph("To ensure that corrupted or illogical vehicle data cannot enter the simulation, the Car class uses System.ComponentModel.DataAnnotations.")
doc.add_paragraph("• Properties such as OperationalRangeKm and BatteryLevelPercent use [Range] attributes to enforce physical limits.", style='List Bullet')
doc.add_paragraph("• The CarId and ModelName properties are marked [Required].", style='List Bullet')
doc.add_paragraph("• The Validate() method ensures data is checked at the time of instantiation and after any manual updates from the control panel. If a rule is violated, a CarDataValidationException is thrown and gracefully caught by the UI.", style='List Bullet')

add_heading("5. Exception Handling and Incident Response")
doc.add_paragraph("Robust error handling is critical for mission-control systems. All simulated incidents inherit from the SimulationException base class, which carries a SeverityLevel and a timestamp.")
doc.add_paragraph("• A central SimulateIncident method in the FleetManager uses a unified try-catch block.", style='List Bullet')
doc.add_paragraph("• When a CriticalBatteryException occurs, the vehicle's state is safely transitioned to Charging.", style='List Bullet')
doc.add_paragraph("• When a NavigationErrorException or UnauthorizedOverrideException occurs, the vehicle is forced to Stopped.", style='List Bullet')
doc.add_paragraph("• This ensures the simulation itself never crashes, but instead resolves incidents into safe states and alerts the operators.", style='List Bullet')

add_heading("6. Delegate-Based Event Notification")
doc.add_paragraph("The simulation logic must remain decoupled from the Windows Forms UI. To achieve this, the AlertSystem relies on C# delegates.")
doc.add_paragraph("• An AlertDispatcher publishes OnAlert events carrying a ControlRoomAlertEventArgs payload.", style='List Bullet')
doc.add_paragraph("• The MainForm subscribes to this event. When an incident occurs (e.g., entering a restricted zone), the FleetManager tells the dispatcher to raise an alert, which asynchronously populates the Mission Control alert history list without the simulation knowing about the UI components.", style='List Bullet')

add_heading("7. LINQ and Dynamic Vehicle Data Management")
doc.add_paragraph("The FleetManager exposes several methods to dynamically query the live fleet data using Language Integrated Query (LINQ):")
doc.add_paragraph("• Filtering: GetCarsByStatus() filters vehicles by their operational state.", style='List Bullet')
doc.add_paragraph("• Sorting: GetPriorityMissions() chains Where, OrderByDescending, and ThenBy to find critical missions with low batteries.", style='List Bullet')
doc.add_paragraph("• Aggregation: GetFleetSummary() uses grouping and averaging to provide a real-time dashboard of the fleet's average battery, total vehicles, and mission priorities.", style='List Bullet')

add_heading("8. User-Centric Control Panel (GUI)")
doc.add_paragraph("The graphical interface is built using standard WinForms controls arranged in a responsive MainForm.")
doc.add_paragraph("• Tab Navigation: The dashboard separates concerns into Fleet (grid view), Alert history (mission events), and Operational logs (per-vehicle history).", style='List Bullet')
doc.add_paragraph("• Control Actions: A dedicated right-hand side panel provides quick controls to add, update, and remove vehicles, as well as trigger dynamic simulation events like Critical Battery and Restricted Zone.", style='List Bullet')
doc.add_paragraph("• Real-time Feedback: The DataGridView and summary labels are automatically refreshed whenever the fleet state changes.", style='List Bullet')
p_holder1 = doc.add_paragraph("[PLACE SCREENSHOT HERE: Main Dashboard]\n(Insert a screenshot showing the main GUI with the populated DataGridView and summary labels.)")
p_holder1.runs[0].font.italic = True
p_holder1.runs[0].font.color.rgb = docx.shared.RGBColor(128, 128, 128)

add_heading("9. Operational Safety, Reliability and Design Justification")
doc.add_paragraph("The application is designed around safety and reliability. By using a strict domain model (Car), we guarantee data integrity. By centralizing exception handling within the FleetManager, we guarantee that no incident goes unhandled. The event-driven architecture ensures that the underlying simulation could theoretically run headlessly or be connected to a different user interface (e.g., WPF or a web API) without modification.")

add_heading("10. Testing Scenarios and Expected Results")
doc.add_paragraph("1. Adding Valid/Invalid Vehicles: Adding a car with 105% battery will trigger a validation error, preventing addition.")
doc.add_paragraph("[PLACE SCREENSHOT HERE: Validation Error MessageBox]", style='Intense Quote')

doc.add_paragraph("2. Executing an Incident: Clicking 'Navigation error' will halt the selected car, append a warning to its operational log, update the fleet grid, and broadcast a critical alert to the mission control panel.")
doc.add_paragraph("[PLACE SCREENSHOT HERE: Alert History Tab]", style='Intense Quote')

doc.add_paragraph("3. Filtering: Selecting a specific operational status from the filter drop-down will immediately update the DataGridView and recalculate the fleet averages via LINQ.")
doc.add_paragraph("4. Log Retention: Selecting a car from the grid will instantly populate the 'Operational logs' tab with that specific car's timestamped history.")
doc.add_paragraph("[PLACE SCREENSHOT HERE: Operational Logs Tab]", style='Intense Quote')

add_heading("11. Data Persistence (JSON)")
doc.add_paragraph("To ensure the simulation state is not lost between sessions, robust data persistence was implemented using the native System.Text.Json library.")
doc.add_paragraph("• Serialization/Deserialization: The FleetManager handles writing the active fleet state to a fleet.json file whenever the application is closed.", style='List Bullet')
doc.add_paragraph("• State Restoration: Upon launching the application, the MainForm automatically checks for fleet.json. If it exists, the fleet, including each car's operational log and current status, is instantly restored. If not, the application falls back to seeding default demo data.", style='List Bullet')
doc.add_paragraph("• JSON Integrity: The Car model's properties are fully compatible with JSON serialization, ensuring seamless translation from memory to disk.", style='List Bullet')

add_heading("12. Conclusion")
doc.add_paragraph("The implemented Self-Driving Car Simulation successfully meets all requirements of the CSC 439 assignment. It demonstrates a robust grasp of object-oriented C# principles, including data annotations, custom exception hierarchies, delegate-driven event notifications, dynamic LINQ queries, and JSON data persistence, all wrapped within a responsive WinForms graphical user interface.")

add_heading("Appendix A. Core Class Implementation Overview")
doc.add_paragraph("(Note: Full source code files are available in the project repository. Key snippets are highlighted here.)")

p = doc.add_paragraph()
r = p.add_run("The Car Model (Data Annotation and Logging):")
r.bold = True
code1 = '''public class Car
{
    [Required]
    public string CarId { get; }
    
    [Range(0, 100)]
    public double BatteryLevelPercent { get; set; }
    
    // Validates object constraints programmatically
    public void Validate()
    {
        var validationContext = new ValidationContext(this);
        var validationResults = new List<ValidationResult>();
        if (!Validator.TryValidateObject(this, validationContext, validationResults, validateAllProperties: true))
            throw new CarDataValidationException(validationResults.First().ErrorMessage);
    }
}'''
doc.add_paragraph(code1)

p2 = doc.add_paragraph()
r2 = p2.add_run("\nFleet Manager (LINQ Aggregation):")
r2.bold = True
code2 = '''public FleetSummary GetFleetSummary()
{
    var cars = _fleet.Values.ToList();
    return new FleetSummary
    {
        TotalCars = cars.Count,
        AverageBatteryPercent = cars.Count == 0 ? 0 : cars.Average(c => c.BatteryLevelPercent),
        LowBatteryCount = cars.Count(c => c.BatteryLevelPercent <= LowBatteryThreshold)
    };
}'''
doc.add_paragraph(code2)

p3 = doc.add_paragraph()
r3 = p3.add_run("\nException Handling (Safe Incident Resolution):")
r3.bold = True
code3 = '''private void HandleSimulationException(string carId, SimulationException ex)
{
    if (_fleet.TryGetValue(carId, out var car))
    {
        car.LogEvent($"INCIDENT [{ex.Severity}]: {ex.Message}");
        if (ex is CriticalBatteryException)
            car.Status = OperationalStatus.Charging;
        else if (ex is NavigationErrorException or UnauthorizedOverrideException)
            car.Status = OperationalStatus.Stopped;
    }
    Alerts.RaiseAlert(carId, ControlRoomEventType.IncidentRaised, ex.Message, ex.Severity);
}'''
doc.add_paragraph(code3)

doc.add_page_break()
add_heading("Appendix B. Application Screenshots")
doc.add_paragraph("(This section is reserved for additional high-resolution screenshots of the application running, to satisfy the documentation requirements.)")
doc.add_paragraph("[PLACE ADDITIONAL SCREENSHOTS HERE]")
doc.add_paragraph("• Figure 1: Main Dashboard displaying active autonomous fleet.", style='List Bullet')
doc.add_paragraph("• Figure 2: Real-time Mission Control Alerts triggered by custom SimulationExceptions.", style='List Bullet')
doc.add_paragraph("• Figure 3: Individual vehicle operational log view.", style='List Bullet')

doc.save('Updated_Assignment_Report.docx')
